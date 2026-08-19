"""Create the background-subtraction task sheet from the retained DOCX template."""

from __future__ import annotations

import hashlib
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REFERENCE = PROJECT_ROOT / "data" / "raw" / "our_dataset" / "科创任务安排.docx"
OUTPUT = PROJECT_ROOT / "docs" / "科创任务安排_背景扣除.docx"
EXPECTED_SHA256 = "8B5877CE79DBF072051FB4F1EED08704FABDDFABD496DD28E0E2888BE1076AD8"


REPLACEMENTS = {
    0: "红外图像二维背景扣除",
    1: "任务背景",
    2: "参考CEERS Epoch 1 NIRCam方法：先压平大尺度起伏，经多层源检测和膨胀构造重源掩膜，再在未掩膜区拟合稳健二维背景。本任务将该思路适配到RKZ50单帧红外科学图像。",
    3: "任务目标",
    4: "对data/processed/flicker输出的160帧1/f校正科学图进行二维背景估计与扣除，排除条纹模型文件，在保护恒星流量和像素噪声的前提下降低大尺度背景不均匀性。",
    5: "主要工作",
    6: "只读取data/processed/flicker中的flicker_corrected_*.fits，排除flicker_model_*.fits。",
    7: "合并blindmap、边缘、无效像元和已知目标星掩膜。",
    8: "用Sigma裁剪和稳健位置估计得到粗二维背景。",
    9: "屏蔽高于5倍稳健RMS的亮像元并估计大环尺度背景。",
    10: "由扩展源到紧致源进行多尺度检测，逐层膨胀源掩膜。",
    11: "在原图未掩膜区域估计最终稳健二维网格背景。",
    12: "中值滤波粗网格并插值为全分辨率背景模型。",
    13: "扣除背景并比较背景分布、尺度相关RMS和高频噪声。",
    14: "检查源掩膜边缘残差和目标星孔径流量。",
    15: "输出产品",
    16: "background_subtracted_*.fits",
    17: "background_model_*.fits",
    18: "background_statistics.csv",
    19: "粗背景、环形背景和四层源掩膜诊断图",
    20: "扣除前后图像、背景分布及尺度相关RMS图",
    21: "全帧目标星测光变化图",
    22: "验收要求",
    23: "两组各80帧均完成处理，且不使用1/f条纹模型FITS。",
    24: "blindmap、源、边缘及无效像元不参与背景估计。",
    25: "64像素尺度背景位置散布下降目标不低于80%。",
    26: "扣除后相邻像元差分噪声不得增加超过2%。",
    27: "SNR不低于10的目标星孔径流量变化不超过1%。",
    28: "质量门失败时输出原图float32副本和零背景模型。",
    29: "输出满足background_subtracted = input - background_model。",
    30: "保留FITS科学头和已有FLK元数据，全部产品通过严格校验。",
    31: "至少包含8个自动测试，覆盖算法、输入选择、测光及FITS输出。",
}


def main() -> None:
    digest = hashlib.sha256(REFERENCE.read_bytes()).hexdigest().upper()
    if digest != EXPECTED_SHA256:
        raise RuntimeError(f"Template hash changed: {digest}")

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    with ZipFile(REFERENCE, "r") as source:
        document_xml = source.read("word/document.xml")
        root = etree.fromstring(document_xml)
        paragraphs = root.xpath(".//w:body/w:p", namespaces=namespace)
        if len(paragraphs) != 32:
            raise RuntimeError(f"Unexpected paragraph count: {len(paragraphs)}")
        for index, text in REPLACEMENTS.items():
            text_nodes = paragraphs[index].xpath(".//w:t", namespaces=namespace)
            if not text_nodes:
                raise RuntimeError(f"Paragraph {index} contains no text node")
            text_nodes[0].text = text
            text_nodes[0].set(xml_space, "preserve")
            for node in text_nodes[1:]:
                node.text = ""
        patched_xml = etree.tostring(
            root, encoding="UTF-8", xml_declaration=True, standalone=True
        )
        with ZipFile(OUTPUT, "w", compression=ZIP_DEFLATED) as destination:
            for info in source.infolist():
                payload = patched_xml if info.filename == "word/document.xml" else source.read(info.filename)
                destination.writestr(info, payload)

    check = Document(OUTPUT)
    expected_numbering = {**{i: 2 for i in range(6, 15)}, **{i: 3 for i in range(16, 22)}, **{i: 4 for i in range(23, 32)}}
    for index, num_id in expected_numbering.items():
        num_pr = check.paragraphs[index]._p.pPr.numPr
        if num_pr is None or int(num_pr.numId.val) != num_id:
            raise RuntimeError(f"Numbering changed at paragraph {index}")
    if check.paragraphs[0].alignment is None or not check.paragraphs[0].runs[0].bold:
        raise RuntimeError("Title formatting was not preserved")
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
